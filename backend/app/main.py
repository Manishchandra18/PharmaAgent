import time
import traceback
import asyncio
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel
from typing import List, Dict, Any
import os
from dotenv import load_dotenv
import fitz  # PyMuPDF for PDF processing
from langchain_perplexity import ChatPerplexity
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage
import tiktoken
from typing import List
from PIL import Image
import io
import pytesseract
import boto3
from botocore.exceptions import ClientError, NoCredentialsError

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Load environment variables
load_dotenv()

# Initialize AWS Textract client
# Use us-east-1 to match Medical Comprehend region
try:
    textract_client = boto3.client(
        'textract',
        region_name='us-east-1',  # Use us-east-1 for consistency with Medical Comprehend
        aws_access_key_id=os.environ.get('AWS_ACCESS_KEY_ID'),
        aws_secret_access_key=os.environ.get('AWS_SECRET_ACCESS_KEY')
    )
    print("[DEBUG] AWS Textract client initialized successfully (us-east-1)")
except Exception as e:
    print(f"[WARNING] AWS Textract client initialization failed: {e}")
    textract_client = None

# Initialize AWS Medical Comprehend client
# Note: Medical Comprehend is only available in specific regions
# Supported regions: us-east-1, us-east-2, us-west-2, eu-west-1, eu-west-2, ap-southeast-2, ca-central-1
try:
    # Use us-east-1 for Medical Comprehend (most reliable region)
    medical_comprehend_client = boto3.client(
        'comprehendmedical',
        region_name='us-east-1',  # Force us-east-1 for Medical Comprehend
        aws_access_key_id=os.environ.get('AWS_ACCESS_KEY_ID'),
        aws_secret_access_key=os.environ.get('AWS_SECRET_ACCESS_KEY')
    )
    print("[DEBUG] AWS Medical Comprehend client initialized successfully (us-east-1)")
except Exception as e:
    print(f"[WARNING] AWS Medical Comprehend client initialization failed: {e}")
    print("[INFO] Medical Comprehend is only available in: us-east-1, us-east-2, us-west-2, eu-west-1, eu-west-2, ap-southeast-2, ca-central-1")
    medical_comprehend_client = None


# Universal Medical Summary Generation Prompt
MEDICAL_SUMMARY_PROMPT = """
You are an expert medical AI assistant specializing in analyzing medical documents and prescriptions. 
Generate ONLY a comprehensive medical summary. Do not include any reasoning, analysis, or commentary. 
Start immediately with "MEDICAL SUMMARY" and follow the exact template below.

CRITICAL INSTRUCTIONS TO THE MODEL:
A. Extract ALL information from the provided documents - do not miss any details
B. Pay special attention to handwritten text, prescription details, and medical notes
C. Include specific dosages, frequencies, durations, and medication names exactly as written
D. Extract all numerical values (lab results, vital signs, measurements) precisely
E. If critical data are absent, write "Not documented in provided records." 
F. Preserve professional medical language, units and abbreviations exactly as written
G. Follow the section order EXACTLY. Do not add, rename or reorder headers. 
H. Populate bracketed fields with structured data from the documents
I. Include ALL medications, dosages, and instructions found in prescriptions
J. Extract patient demographics, contact info, and insurance details if available
K. Include all diagnostic tests, results, and clinical findings mentioned
L. Note any follow-up appointments, referrals, or special instructions
M. If information is unclear or partially visible, note it as "unclear" rather than omitting it
N. NEVER fabricate or assume values not explicitly stated in the documents


=======================================================================
                        MEDICAL SUMMARY
=======================================================================

### HEADER SECTION
Hospital Details
 • [Hospital Name & Letterhead]
 • Department/Ward: [e.g., Neuro ICU / Orthopaedics Ward]

Patient Details
 • Name: [Full legal name]
 • UHID / MRN: [Unique hospital ID]
 • Age / Gender: [Calculated age; sex]
 • Date of Birth: [DD-MM-YYYY]
 • Address: [Full postal address]
 • Mobile No.: [Primary contact]
 • Admission Date & Time: [DD-MMM-YYYY, HH:MM]
 • Medical / LAMA Date & Time: [DD-MMM-YYYY, HH:MM]
 • IP / Encounter No.: [IP number]
 • Bed & Ward: [e.g., Bed 2304, 2nd-Floor HDU]
Treating Consultant: [Name, Qualification, Specialty]
Additional Consultants: [Specialists consulted, if any]
Next of Kin: [Name, Relationship, Contact Number]

---

### 1. FINAL DIAGNOSIS
Primary Diagnosis:
 – […]
Secondary Diagnoses / Co-morbidities:
 – […]
ICD-10 Codes: […]

### 2. CHIEF COMPLAINT
[Presenting symptoms with duration]

### 3. CLINICAL FINDINGS ON ADMISSION
Vital Signs (initial assessment)
| Parameter        | Value | Unit    | Reference Range | Status |
|------------------|-------|---------|-----------------|--------|
| Temperature      | […]   | °C / °F | […]             | […]    |
| Blood Pressure   | […]   | mmHg    | […]             | […]    |
| Heart Rate       | […]   | bpm     | […]             | […]    |
| Respiratory Rate | […]   | /min    | […]             | […]    |
| SpO₂             | […]   | %       | […]             | […]    |

Physical Examination Highlights
 • General Condition: […]
 • CVS: […]
 • RESP: […]
 • CNS: […]
 • PA: […]
 • Local / Neuro / Ortho Exam (if applicable): […]

### 4. RELEVANT HISTORY
Past Medical History:
 – […]
Past Surgical / Procedural History:
 – […]
Family History:
 – […]
Personal & Social History:
 – […]
Drug / Food Allergies: […]
Current Medications on admission:
 – […]
 – […]

### 5. HOSPITAL COURSE (CHRONOLOGY)
Day-wise narrative covering:
 • Diagnostic work-up & results
 • Therapeutic interventions (medications, surgeries, ventilation, dialysis)
 • Response to treatment, complications, specialist opinions
 • Transfers (ICU ↔ ward) and significant events
Interdepartmental references

### 6. INVESTIGATIONS
Key Laboratory Results (dates & units):
 – […]
Imaging Results:
 – […]
Special Investigations:
 – […]
(Attach full reports separately in EMR.)

### 7. TREATMENTS & PROCEDURES DURING STAY
| Date       | Procedure                                        | Anaesthesia      | Key Findings / Implants           |
|------------|--------------------------------------------------|------------------|-----------------------------------|
| […]        | […]                                              | […]              | […]                               |

Treatments received:
 – […]
 – […]
 – […]

Implants Used (if any):
 – […]

### 8. CONDITION AT MEDICAL / LAMA
Vital Signs
| Parameter        | Value | Unit    | Reference Range | Status |
|------------------|-------|---------|-----------------|--------|
| Temperature      | […]   | °C / °F | […]             | […]    |
| Blood Pressure   | […]   | mmHg    | […]             | […]    |
| Heart Rate       | […]   | bpm     | […]             | […]    |
| Respiratory Rate | […]   | /min    | […]             | […]    |
| SpO₂             | […]   | %       | […]             | […]    |

Clinical Status
 • General: Stable / Improved / Deteriorated / On NIV
 • Neurological: e.g., “GCS 15/15, obeys commands”
 • Wound / Surgical Site: […]
 • Mobility & ADL Status: […]

### 9. MEDICATIONS
| Medicine          | Dose     | Frequency | Route | Duration / Stop Date |
|-------------------|----------|-----------|-------|----------------------|
| […]               | […]      | […]       | […]   | […]                  |

### 10. MEDICAL ADVICE & INSTRUCTIONS
 1. Diet: […]
 2. Activity & Physiotherapy: […]
 3. Wound / Drain / Catheter Care: […]
 4. Device Precautions: […]
 5. Red-Flag / Warning Signs: […]
 6. Lifestyle & Preventive Advice: […]
 7. Medication Compliance: […]

### 11. FOLLOW-UP PLAN
 • Next OPD Visit: [Date, Dept., Consultant]
 • Scheduled Investigations before visit: […]
 • Emergency Contact: [Contact your local emergency services]

### 12. PHYSICIAN ATTESTATION
[Signature] [Full Name, Degrees]
[Designation & Department] [Medical Reg. No.]
Date: [Medical / LAMA date]

### 13. PATIENT ACKNOWLEDGEMENT
“I / we have been explained the medication and precautions required in detail by nursing staff. All documents and belongings were handed over.”

**Patient/Relative Signature:** __________  
**Name:** [Next of Kin & Relationship]  
**Date:** [Medical Date]  
**Time:** [Medical Time]

=======================================================================  
                      END OF SUMMARY  
=======================================================================
"""


# Initialize Perplexity Client
api_key = os.environ.get("PERPLEXITY_API_KEY")
print(f"[DEBUG] API Key found: {'Yes' if api_key else 'No'}")
if api_key:
    print(f"[DEBUG] API Key length: {len(api_key)}")
    print(f"[DEBUG] API Key starts with: {api_key[:10]}...")

try:
    perplexity_client = ChatPerplexity(
        model="sonar-pro",
        temperature=0.2,
        pplx_api_key=api_key,
    )
    print("[DEBUG] Perplexity client initialized successfully")
except Exception as e:
    print(f"[ERROR] initializing Perplexity client: {e}")
    perplexity_client = None


app = FastAPI(title="Medical AI Agent API", version="0.1.0", root_path="/api/v1")


app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class SummaryRequest(BaseModel):
    context: str


class ChatMessage(BaseModel):
    role: str
    content: str


class ChatRequest(BaseModel):
    messages: List[ChatMessage]




tokenizer = tiktoken.encoding_for_model("gpt-4")


def estimate_input_tokens(text: str) -> int:
    return len(tokenizer.encode(text))


def calculate_max_tokens(count: int) -> int:
    if count < 5_000:
        return 4_000
    elif count < 10_000:
        return 7_000
    elif count < 20_000:
        return 10_000
    return 12_000


# Thread-pool PDF extraction
from concurrent.futures import ThreadPoolExecutor
executor = ThreadPoolExecutor()


from PIL import Image


def calculate_confidence_score(text: str) -> float:
    """Calculate confidence score based on text characteristics"""
    if not text or len(text.strip()) < 10:
        return 0.0
    
    score = 0.0
    text = text.strip()
    
    # Length factor (longer text generally more reliable)
    length_score = min(len(text) / 1000, 1.0) * 0.2
    score += length_score
    
    # Word count factor
    words = text.split()
    word_score = min(len(words) / 100, 1.0) * 0.2
    score += word_score
    
    # Character diversity factor (more diverse = likely real text)
    unique_chars = len(set(text.lower()))
    char_diversity = min(unique_chars / 30, 1.0) * 0.2
    score += char_diversity
    
    # Medical terminology presence (basic check)
    medical_terms = ['patient', 'diagnosis', 'treatment', 'medication', 'hospital', 'doctor', 'medical', 'condition', 'symptoms', 'history']
    medical_count = sum(1 for term in medical_terms if term.lower() in text.lower())
    medical_score = min(medical_count / 5, 1.0) * 0.2
    score += medical_score
    
    # Punctuation and structure factor
    punct_count = sum(1 for c in text if c in '.,;:!?')
    punct_score = min(punct_count / 20, 1.0) * 0.2
    score += punct_score
    
    return min(score, 1.0)

def post_process_text(text: str) -> str:
    """Post-process extracted text to improve quality"""
    if not text:
        return text
    
    # Remove excessive whitespace
    import re
    text = re.sub(r'\s+', ' ', text)
    
    # Fix common OCR errors
    replacements = {
        '0': 'O',  # Common OCR confusion
        '1': 'I',  # In certain contexts
        '5': 'S',  # In certain contexts
        '8': 'B',  # In certain contexts
    }
    
    # Only apply replacements in specific contexts to avoid over-correction
    for old, new in replacements.items():
        # Only replace if it's clearly a character (surrounded by letters)
        text = re.sub(f'(?<=[a-zA-Z]){old}(?=[a-zA-Z])', new, text)
    
    return text.strip()

def extract_medical_fields(text: str) -> dict:
    """Extract basic medical fields from text for debugging purposes"""
    if not text:
        return {}
    
    import re
    
    # Convert to lowercase for case-insensitive matching
    text_lower = text.lower()
    
    # Extract common medical patterns
    fields = {}
    
    # Patient ID patterns
    patient_id_patterns = [
        r'uhid[:\s]*([a-zA-Z0-9]+)',
        r'mrn[:\s]*([a-zA-Z0-9]+)',
        r'patient\s*id[:\s]*([a-zA-Z0-9]+)',
        r'ip\s*no[:\s]*([a-zA-Z0-9]+)'
    ]
    
    for pattern in patient_id_patterns:
        match = re.search(pattern, text_lower)
        if match:
            fields['patient_id'] = match.group(1)
            break
    
    # Age patterns
    age_match = re.search(r'age[:\s]*(\d+)', text_lower)
    if age_match:
        fields['age'] = age_match.group(1)
    
    # Gender patterns
    gender_match = re.search(r'(male|female|m|f)[\s,.]', text_lower)
    if gender_match:
        fields['gender'] = gender_match.group(1)
    
    # Date patterns
    date_patterns = [
        r'(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',  # DD/MM/YYYY or DD-MM-YYYY
        r'(\d{1,2}\s+\w+\s+\d{4})'  # DD Month YYYY
    ]
    
    dates = []
    for pattern in date_patterns:
        matches = re.findall(pattern, text_lower)
        dates.extend(matches)
    
    if dates:
        fields['dates_found'] = dates[:3]  # Limit to first 3 dates
    
    # Medical terms count
    medical_terms = [
        'diagnosis', 'treatment', 'medication', 'prescription', 'hospital',
        'patient', 'doctor', 'medical', 'condition', 'symptoms', 'history',
        'blood', 'pressure', 'temperature', 'heart', 'rate', 'respiratory'
    ]
    
    medical_term_count = sum(1 for term in medical_terms if term in text_lower)
    fields['medical_terms_count'] = medical_term_count
    
    # Medication patterns
    medication_patterns = [
        r'(\w+\s*\d+\s*mg)',  # Drug name with dosage
        r'(tablet|capsule|injection|syrup)\s+(\w+)',  # Medication forms
        r'(\w+)\s+(once|twice|thrice|daily|bid|tid|qid)',  # Drug with frequency
    ]
    
    medications = []
    for pattern in medication_patterns:
        matches = re.findall(pattern, text_lower)
        medications.extend(matches)
    
    if medications:
        fields['medications_found'] = medications[:5]  # Limit to first 5
    
    # Vital signs patterns
    vital_signs = {}
    
    # Blood pressure
    bp_match = re.search(r'bp[:\s]*(\d+/\d+)', text_lower)
    if bp_match:
        vital_signs['blood_pressure'] = bp_match.group(1)
    
    # Temperature
    temp_match = re.search(r'temp[:\s]*(\d+\.?\d*)\s*[°]?[cf]?', text_lower)
    if temp_match:
        vital_signs['temperature'] = temp_match.group(1)
    
    # Heart rate
    hr_match = re.search(r'(hr|heart\s*rate)[:\s]*(\d+)', text_lower)
    if hr_match:
        vital_signs['heart_rate'] = hr_match.group(2)
    
    if vital_signs:
        fields['vital_signs'] = vital_signs
    
    # Text statistics
    fields['text_stats'] = {
        'total_characters': len(text),
        'total_words': len(text.split()),
        'total_lines': len(text.split('\n')),
        'has_numbers': bool(re.search(r'\d', text)),
        'has_uppercase': bool(re.search(r'[A-Z]', text)),
        'has_punctuation': bool(re.search(r'[.,;:!?]', text))
    }
    
    return fields

def validate_document_format(content_type: str, file_bytes: bytes) -> tuple[bool, str]:
    """Validate document format and provide helpful error messages"""
    
    # Check content type
    supported_types = ["application/pdf", "image/jpeg", "image/png", "image/tiff"]
    if content_type not in supported_types:
        return False, f"Unsupported content type: {content_type}. Supported formats: PDF, JPEG, PNG, TIFF"
    
    # Validate file signatures
    if content_type == "application/pdf":
        if not file_bytes.startswith(b'%PDF'):
            return False, "Invalid PDF file: File does not start with PDF signature"
        if len(file_bytes) < 100:
            return False, "Invalid PDF file: File too small to be a valid PDF"
            
    elif content_type == "image/jpeg":
        if not file_bytes.startswith(b'\xff\xd8\xff'):
            return False, "Invalid JPEG file: File does not start with JPEG signature"
            
    elif content_type == "image/png":
        if not file_bytes.startswith(b'\x89PNG'):
            return False, "Invalid PNG file: File does not start with PNG signature"
            
    elif content_type == "image/tiff":
        if not (file_bytes.startswith(b'II*\x00') or file_bytes.startswith(b'MM\x00*')):
            return False, "Invalid TIFF file: File does not start with TIFF signature"
    
    # Check file size limits (AWS Textract has limits)
    max_size = 10 * 1024 * 1024  # 10MB
    if len(file_bytes) > max_size:
        return False, f"File too large: {len(file_bytes)} bytes. Maximum size: {max_size} bytes"
    
    if len(file_bytes) < 100:
        return False, "File too small: File appears to be empty or corrupted"
    
    return True, "Valid"

def get_detailed_error_message(error: Exception, content_type: str, filename: str = None) -> str:
    """Generate detailed error messages for better debugging"""
    
    error_str = str(error).lower()
    
    if "unsupporteddocumentexception" in error_str:
        if content_type == "application/pdf":
            return f"PDF processing error: AWS Textract requires PDF files to be processed asynchronously. The current implementation converts PDFs to images for processing. Please ensure the PDF is not password-protected or corrupted."
        elif content_type == "image/tiff":
            return f"TIFF processing error: AWS Textract requires TIFF files to be processed asynchronously. The current implementation converts TIFFs to PNG for processing."
        else:
            return f"Document format error: {content_type} is not supported by AWS Textract's synchronous API. Supported formats: PNG, JPEG (for sync), PDF, TIFF (for async)."
    
    elif "invalidparameter" in error_str:
        return f"Invalid parameters: The document may be corrupted or in an unsupported format. Please check the file and try again."
    
    elif "accessdenied" in error_str or "unauthorized" in error_str:
        return "AWS access error: Please check your AWS credentials and permissions for Textract service."
    
    elif "throttling" in error_str or "rate" in error_str:
        return "Rate limit exceeded: Too many requests to AWS Textract. Please wait a moment and try again."
    
    elif "timeout" in error_str:
        return "Request timeout: The document processing took too long. Try with a smaller file or check your internet connection."
    
    elif "network" in error_str or "connection" in error_str:
        return "Network error: Unable to connect to AWS Textract. Please check your internet connection."
    
    else:
        # Generic error with more context
        base_msg = f"Document processing failed for {content_type}"
        if filename:
            base_msg += f" (file: {filename})"
        base_msg += f": {str(error)}"
        return base_msg

def analyze_medical_text_with_comprehend(text: str) -> dict:
    """Analyze medical text using AWS Medical Comprehend"""
    if not medical_comprehend_client or not text:
        return {"entities": [], "relationships": [], "phi": []}
    
    try:
        print(f"[DEBUG] Analyzing medical text with Comprehend Medical ({len(text)} characters)")
        
        # Detect medical entities
        entities_response = medical_comprehend_client.detect_entities_v2(Text=text)
        entities = entities_response.get('Entities', [])
        
        # Detect relationships between entities
        relationships_response = medical_comprehend_client.infer_icd10_cm(Text=text)
        relationships = relationships_response.get('Entities', [])
        
        # Detect PHI (Protected Health Information)
        phi_response = medical_comprehend_client.detect_phi(Text=text)
        phi = phi_response.get('Entities', [])
        
        print(f"[DEBUG] Medical Comprehend analysis complete: {len(entities)} entities, {len(relationships)} relationships, {len(phi)} PHI items")
        
        return {
            "entities": entities,
            "relationships": relationships,
            "phi": phi,
            "medical_insights": {
                "conditions": [e for e in entities if e.get('Category') == 'MEDICAL_CONDITION'],
                "medications": [e for e in entities if e.get('Category') == 'MEDICATION'],
                "dosages": [attr for e in entities if e.get('Category') == 'MEDICATION' for attr in e.get('Attributes', []) if attr.get('Type') == 'DOSAGE'],
                "procedures": [e for e in entities if e.get('Category') == 'PROCEDURE'],
                "anatomy": [e for e in entities if e.get('Category') == 'SYSTEM_ORGAN_SITE'],
            }
        }
        
    except Exception as e:
        print(f"[ERROR] Medical Comprehend analysis failed: {e}")
        return {"entities": [], "relationships": [], "phi": [], "error": str(e)}

def extract_text_with_aws_textract_sync(image_bytes: bytes, content_type: str) -> tuple[str, float]:
    """Extract text using AWS Textract synchronous API (PNG/JPEG only)"""
    if not textract_client:
        print("[DEBUG] AWS Textract client not available - this should not happen!")
        raise Exception("AWS Textract client not available")
    
    print(f"[DEBUG] Starting AWS Textract sync extraction on {len(image_bytes)} bytes, content_type: {content_type}")
    
    # Validate format for sync API
    if content_type not in ['image/png', 'image/jpeg']:
        raise Exception(f"Unsupported document format for sync processing: {content_type}. Use PNG or JPEG for synchronous processing.")
    
    try:
        # Call AWS Textract analyze_document for better form/table detection
        print("[DEBUG] Calling AWS Textract analyze_document with FORMS and TABLES features")
        response = textract_client.analyze_document(
            Document={'Bytes': image_bytes},
            FeatureTypes=['FORMS', 'TABLES']
        )
        
        print(f"[DEBUG] AWS Textract response received with {len(response.get('Blocks', []))} blocks")
        
        # Extract text and confidence scores
        extracted_text = ""
        confidence_scores = []
        line_count = 0
        
        for block in response.get('Blocks', []):
            if block['BlockType'] == 'LINE':
                text = block.get('Text', '')
                confidence = block.get('Confidence', 0)
                
                if text.strip():
                    extracted_text += text + "\n"
                    confidence_scores.append(confidence)
                    line_count += 1
                    print(f"[DEBUG] Line {line_count}: '{text.strip()}' (confidence: {confidence:.1f}%)")
        
        # Calculate average confidence
        avg_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0.0
        
        print(f"[DEBUG] AWS Textract extraction complete: {line_count} lines, avg confidence: {avg_confidence:.1f}%")
        print(f"[DEBUG] Extracted text preview: {extracted_text[:200]}...")
        
        return extracted_text.strip(), avg_confidence / 100.0  # Convert to 0-1 scale
        
    except ClientError as e:
        error_code = e.response.get('Error', {}).get('Code', 'Unknown')
        error_message = e.response.get('Error', {}).get('Message', str(e))
        print(f"[ERROR] AWS Textract ClientError: {error_code} - {error_message}")
        
        if error_code == 'UnsupportedDocumentException':
            raise Exception(f"Document format not supported by AWS Textract sync API. Supported formats: PNG, JPEG. Received: {content_type}")
        else:
            raise Exception(f"AWS Textract failed ({error_code}): {error_message}")
    except Exception as e:
        print(f"[ERROR] Unexpected error with AWS Textract: {e}")
        raise Exception(f"AWS Textract failed: {e}")

def extract_text_with_aws_textract_async(document_bytes: bytes, content_type: str) -> tuple[str, float]:
    """Extract text using AWS Textract asynchronous API (PDF/TIFF)"""
    if not textract_client:
        print("[DEBUG] AWS Textract client not available - this should not happen!")
        raise Exception("AWS Textract client not available")
    
    print(f"[DEBUG] Starting AWS Textract async extraction on {len(document_bytes)} bytes, content_type: {content_type}")
    
    # For now, we'll use a simplified approach for async processing
    # In a production environment, you'd want to implement proper S3 integration
    # and job polling, but for this demo we'll use a fallback approach
    
    try:
        # For PDFs, we'll convert to images and process synchronously
        if content_type == "application/pdf":
            print("[DEBUG] PDF detected - using image conversion approach")
            return extract_text_from_pdf_via_images(document_bytes)
        elif content_type == "image/tiff":
            print("[DEBUG] TIFF detected - using image conversion approach") 
            return extract_text_from_tiff_via_images(document_bytes)
        else:
            raise Exception(f"Unsupported format for async processing: {content_type}")
            
    except Exception as e:
        print(f"[ERROR] AWS Textract async processing failed: {e}")
        raise Exception(f"AWS Textract async processing failed: {e}")

def extract_text_from_pdf_via_images(pdf_bytes: bytes) -> tuple[str, float]:
    """Extract text from PDF by converting to images and processing with sync API"""
    try:
        import fitz  # PyMuPDF
        
        # Open PDF
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        all_text = ""
        all_confidences = []
        
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            
            # Convert page to image
            mat = fitz.Matrix(2.0, 2.0)  # Scale factor for better quality
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            
            # Process with sync API
            try:
                page_text, confidence = extract_text_with_aws_textract_sync(img_data, "image/png")
                all_text += page_text + "\n\n"
                all_confidences.append(confidence)
                print(f"[DEBUG] Processed PDF page {page_num + 1}/{pdf_document.page_count}")
            except Exception as e:
                print(f"[WARNING] Failed to process PDF page {page_num + 1}: {e}")
                # Fallback to PyMuPDF text extraction
                page_text = page.get_text()
                all_text += page_text + "\n\n"
                all_confidences.append(0.5)  # Lower confidence for fallback
        
        pdf_document.close()
        
        avg_confidence = sum(all_confidences) / len(all_confidences) if all_confidences else 0.0
        print(f"[DEBUG] PDF processing complete: {len(all_text)} characters, avg confidence: {avg_confidence:.2f}")
        
        return all_text.strip(), avg_confidence
        
    except ImportError:
        print("[ERROR] PyMuPDF not available for PDF processing")
        raise Exception("PDF processing requires PyMuPDF library")
    except Exception as e:
        print(f"[ERROR] PDF processing failed: {e}")
        raise Exception(f"PDF processing failed: {e}")

def extract_text_from_tiff_via_images(tiff_bytes: bytes) -> tuple[str, float]:
    """Extract text from TIFF by converting to PNG and processing with sync API"""
    try:
        from PIL import Image
        
        # Open TIFF image
        tiff_image = Image.open(io.BytesIO(tiff_bytes))
        all_text = ""
        all_confidences = []
        
        # Handle multi-page TIFFs
        page_num = 0
        try:
            while True:
                # Convert to PNG
                png_buffer = io.BytesIO()
                tiff_image.save(png_buffer, format='PNG')
                png_data = png_buffer.getvalue()
                
                # Process with sync API
                try:
                    page_text, confidence = extract_text_with_aws_textract_sync(png_data, "image/png")
                    all_text += page_text + "\n\n"
                    all_confidences.append(confidence)
                    page_num += 1
                    print(f"[DEBUG] Processed TIFF page {page_num}")
                except Exception as e:
                    print(f"[WARNING] Failed to process TIFF page {page_num}: {e}")
                    all_confidences.append(0.5)  # Lower confidence for failed page
                
                # Move to next page
                tiff_image.seek(tiff_image.tell() + 1)
                
        except EOFError:
            # No more pages
            pass
        
        # If single page TIFF, process it
        if page_num == 0:
            png_buffer = io.BytesIO()
            tiff_image.save(png_buffer, format='PNG')
            png_data = png_buffer.getvalue()
            
            try:
                page_text, confidence = extract_text_with_aws_textract_sync(png_data, "image/png")
                all_text += page_text
                all_confidences.append(confidence)
                page_num = 1
            except Exception as e:
                print(f"[WARNING] Failed to process TIFF: {e}")
                all_confidences.append(0.5)
        
        avg_confidence = sum(all_confidences) / len(all_confidences) if all_confidences else 0.0
        print(f"[DEBUG] TIFF processing complete: {len(all_text)} characters, avg confidence: {avg_confidence:.2f}")
        
        return all_text.strip(), avg_confidence
        
    except Exception as e:
        print(f"[ERROR] TIFF processing failed: {e}")
        raise Exception(f"TIFF processing failed: {e}")

def extract_text_with_aws_textract(image_bytes: bytes, content_type: str = None) -> tuple[str, float]:
    """Extract text using AWS Textract - automatically choose sync or async based on format"""
    if not content_type:
        # Try to detect content type from bytes
        if image_bytes.startswith(b'%PDF'):
            content_type = 'application/pdf'
        elif image_bytes.startswith(b'\xff\xd8\xff'):
            content_type = 'image/jpeg'
        elif image_bytes.startswith(b'\x89PNG'):
            content_type = 'image/png'
        elif image_bytes.startswith(b'II*\x00') or image_bytes.startswith(b'MM\x00*'):
            content_type = 'image/tiff'
        else:
            raise Exception("Unable to detect document format")
    
    print(f"[DEBUG] Detected content type: {content_type}")
    
    # Choose appropriate processing method
    if content_type in ['image/png', 'image/jpeg']:
        return extract_text_with_aws_textract_sync(image_bytes, content_type)
    elif content_type in ['application/pdf', 'image/tiff']:
        return extract_text_with_aws_textract_async(image_bytes, content_type)
    else:
        raise Exception(f"Unsupported document format: {content_type}")


def extract_text_with_aws_textract_only(image: Image.Image) -> tuple[str, float]:
    """Extract text using AWS Textract only"""
    try:
        # Convert PIL image to bytes
        img_buffer = io.BytesIO()
        image.save(img_buffer, format='PNG')
        img_bytes = img_buffer.getvalue()
        
        # Use AWS Textract with PNG content type
        text, confidence = extract_text_with_aws_textract(img_bytes, 'image/png')
        
        # Post-process the text
        processed_text = post_process_text(text)
        
        return processed_text, confidence
        
    except Exception as e:
        print(f"[ERROR] AWS Textract extraction failed: {e}")
        return f"[Error: {str(e)}]", 0.0

async def extract_text(pdf_bytes: bytes) -> tuple[str, float]:
    """Extract text from PDF using AWS Textract directly"""
    def _extract(b: bytes):
        try:
            print(f"[DEBUG] Processing PDF with AWS Textract directly ({len(b)} bytes)")
            
            # Use AWS Textract directly on PDF with proper content type
            extracted_text, confidence = extract_text_with_aws_textract(b, 'application/pdf')
            
            print(f"[DEBUG] PDF extraction complete: {len(extracted_text)} characters, confidence: {confidence:.2f}")
            
            return extracted_text, confidence
            
        except Exception as e:
            print(f"[ERROR] PDF extraction failed: {e}")
            return f"[Error: {str(e)}]", 0.0
    
    return await asyncio.get_event_loop().run_in_executor(executor, _extract, pdf_bytes)

async def extract_text_from_image(image_bytes: bytes, content_type: str = None) -> tuple[str, float]:
    """Extract text from image files (JPEG, PNG, TIFF) using AWS Textract"""
    def _extract_from_image(b: bytes, ct: str):
        try:
            # For TIFF, use the new async processing
            if ct == 'image/tiff':
                return extract_text_with_aws_textract(b, ct)
            
            # For PNG/JPEG, convert to PIL image first for better handling
            img = Image.open(io.BytesIO(b))
            
            # Convert to RGB if necessary
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Use AWS Textract with the specific content type
            if ct:
                return extract_text_with_aws_textract(b, ct)
            else:
                return extract_text_with_aws_textract_only(img)
        except Exception as e:
            return f"[Error processing image: {str(e)}]", 0.0
    
    return await asyncio.get_event_loop().run_in_executor(executor, _extract_from_image, image_bytes, content_type)

@app.post("/medical-analysis")
async def get_medical_analysis(request: SummaryRequest):
    """Get detailed medical analysis using AWS Medical Comprehend"""
    if not request.context:
        raise HTTPException(status_code=400, detail="No text provided for analysis")
    
    medical_analysis = analyze_medical_text_with_comprehend(request.context)
    
    return {
        "medical_analysis": medical_analysis,
        "summary": {
            "total_entities": len(medical_analysis.get("entities", [])),
            "medical_conditions": len(medical_analysis.get("medical_insights", {}).get("conditions", [])),
            "medications": len(medical_analysis.get("medical_insights", {}).get("medications", [])),
            "procedures": len(medical_analysis.get("medical_insights", {}).get("procedures", [])),
            "phi_detected": len(medical_analysis.get("phi", []))
        }
    }

@app.post("/debug/extract")
async def debug_extract_text(files: List[UploadFile] = File(...)):
    """Debug endpoint to view raw extracted text without LLM processing"""
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")
    
    results = []
    
    for file in files:
        if file.content_type not in ["application/pdf", "image/jpeg", "image/png", "image/tiff"]:
            continue
        
        try:
            file_bytes = await file.read()
            print(f"[DEBUG] Processing file: {file.filename} ({len(file_bytes)} bytes)")
            
            # Validate document format first
            is_valid, validation_msg = validate_document_format(file.content_type, file_bytes)
            if not is_valid:
                results.append({
                    "filename": file.filename,
                    "error": f"Validation failed: {validation_msg}",
                    "extraction_method": "aws_textract"
                })
                continue
            
            # Handle different file types
            if file.content_type == "application/pdf":
                text, confidence = await extract_text(file_bytes)
            else:
                text, confidence = await extract_text_from_image(file_bytes, file.content_type)
            
            # Extract medical fields for debugging
            medical_fields = extract_medical_fields(text)
            
            results.append({
                "filename": file.filename,
                "content_type": file.content_type,
                "file_size": len(file_bytes),
                "extracted_text": text,
                "confidence": confidence,
                "text_length": len(text),
                "word_count": len(text.split()),
                "medical_fields": medical_fields,
                "extraction_method": "aws_textract"
            })
            
        except Exception as e:
            # Generate detailed error message
            detailed_error = get_detailed_error_message(e, file.content_type, file.filename)
            results.append({
                "filename": file.filename,
                "error": detailed_error,
                "extraction_method": "aws_textract"
            })
    
    return {
        "files_processed": len(results),
        "results": results
    }

@app.post("/upload")
async def upload_files(files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")
    start = time.time()
    combined_text = ""
    extraction_quality = []
    
    for file in files:
        if file.content_type not in ["application/pdf", "image/jpeg", "image/png", "image/tiff"]:
            raise HTTPException(status_code=400, detail=f"'{file.filename}' is not a supported format (PDF, JPEG, PNG, TIFF)")
        
        try:
            file_bytes = await file.read()
            
            # Validate document format first
            is_valid, validation_msg = validate_document_format(file.content_type, file_bytes)
            if not is_valid:
                raise HTTPException(status_code=400, detail=f"Validation failed for {file.filename}: {validation_msg}")
            
            # Handle different file types
            if file.content_type == "application/pdf":
                text, confidence = await extract_text(file_bytes)
            else:
                # Handle image files directly with content type
                text, confidence = await extract_text_from_image(file_bytes, file.content_type)
            
            # Use AWS Textract confidence or fallback to calculated score
            quality_score = confidence if confidence > 0 else calculate_confidence_score(text)
            
            extraction_quality.append({
                "filename": file.filename,
                "quality_score": quality_score,
                "text_length": len(text),
                "word_count": len(text.split()),
                "ocr_method": "aws_textract"
            })
            
            combined_text += text + "\n\n"
        except HTTPException:
            # Re-raise HTTP exceptions as-is
            raise
        except Exception as e:
            # Generate detailed error message
            detailed_error = get_detailed_error_message(e, file.content_type, file.filename)
            raise HTTPException(status_code=500, detail=detailed_error)
    
    elapsed = time.time() - start
    if not combined_text.strip():
        raise HTTPException(status_code=400, detail="No text could be extracted")
    
    # Analyze with Medical Comprehend
    medical_analysis = analyze_medical_text_with_comprehend(combined_text)
    
    return {
        "message": "Files processed successfully",
        "extracted_text": combined_text,
        "time_taken": elapsed,
        "extraction_quality": extraction_quality,
        "overall_quality": sum(q["quality_score"] for q in extraction_quality) / len(extraction_quality) if extraction_quality else 0,
        "medical_analysis": medical_analysis
    }

@app.post("/generate-summary")
async def generate_summary(request: SummaryRequest):
    if not perplexity_client:
        api_key = os.environ.get("PERPLEXITY_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="Perplexity API key not found. Please set PERPLEXITY_API_KEY in your .env file.")
        else:
            raise HTTPException(status_code=500, detail="Perplexity client not initialized. Check your API key.")
    
    # Validate context
    if not request.context or not request.context.strip():
        raise HTTPException(status_code=400, detail="Request context is empty.")

    # Prepare and validate system message
    system_content = MEDICAL_SUMMARY_PROMPT.strip()
    if not system_content:
        raise HTTPException(status_code=500, detail="System prompt is empty.")
    system_message = SystemMessage(content=system_content)

    # Prepare and validate human message
    human_content = request.context.strip()
    if not human_content:
        raise HTTPException(status_code=400, detail="Human message content is empty.")
    human_message = HumanMessage(content=human_content)

    # Debug print raw messages before invocation
    print("[DEBUG] system_message.content length:", len(system_message.content))
    print("[DEBUG] human_message.content length:", len(human_message.content))

    # Estimate tokens
    full_input = system_message.content + "\n\n" + human_message.content
    input_tokens = estimate_input_tokens(full_input)
    max_out = calculate_max_tokens(input_tokens)
    print(f"[DEBUG] tokens in: {input_tokens}, max_out: {max_out}")

    try:
        start = time.time()
        response = await perplexity_client.ainvoke([system_message, human_message])
        elapsed = time.time() - start
        print(f"[DEBUG] summary generation time: {elapsed:.2f}s")
        return {
            "summary": response.content,
            "input_tokens": input_tokens,
            "max_tokens_used": max_out,
            "time_taken": elapsed
        }
    except Exception as e:
        print(f"[ERROR] Perplexity API error: {str(e)}")
        traceback.print_exc()
        
        # Provide more specific error messages
        if "401" in str(e) or "Authorization" in str(e):
            raise HTTPException(status_code=500, detail="Invalid Perplexity API key. Please check your API key in the .env file.")
        elif "429" in str(e):
            raise HTTPException(status_code=500, detail="Rate limit exceeded. Please try again later.")
        else:
            raise HTTPException(status_code=500, detail=f"Summary generation failed: {str(e)}")


@app.post("/chat")
async def chat(request: ChatRequest):
    """General-purpose medical chat endpoint used by the homepage chat widget."""
    if not perplexity_client:
        api_key = os.environ.get("PERPLEXITY_API_KEY")
        if not api_key:
            raise HTTPException(
                status_code=500,
                detail="Perplexity API key not found. Please set PERPLEXITY_API_KEY in your .env file.",
            )
        else:
            raise HTTPException(
                status_code=500,
                detail="Perplexity client not initialized. Check your API key.",
            )

    if not request.messages:
        raise HTTPException(status_code=400, detail="No messages provided.")

    # Build conversation for the LLM
    lc_messages: List = [
        SystemMessage(
            content=(
                "You are a helpful medical AI assistant. "
                "You can explain medical terminology, lab values, diagnoses, and treatment summaries "
                "in clear language. Do not provide medical advice, diagnoses, or treatment plans; "
                "instead, encourage users to consult a licensed healthcare professional for medical decisions."
            )
        )
    ]

    for msg in request.messages:
        content = (msg.content or "").strip()
        if not content:
            continue
        if msg.role == "assistant":
            lc_messages.append(AIMessage(content=content))
        else:
            # Treat anything that is not 'assistant' as a human message
            lc_messages.append(HumanMessage(content=content))

    if len(lc_messages) <= 1:
        raise HTTPException(status_code=400, detail="No valid messages to process.")

    try:
        start = time.time()
        response = await perplexity_client.ainvoke(lc_messages)
        elapsed = time.time() - start
        print(f"[DEBUG] chat generation time: {elapsed:.2f}s")
        return {
            "reply": response.content,
            "time_taken": elapsed,
        }
    except Exception as e:
        print(f"[ERROR] Perplexity chat API error: {str(e)}")
        traceback.print_exc()

        if "401" in str(e) or "Authorization" in str(e):
            raise HTTPException(
                status_code=500,
                detail="Invalid Perplexity API key. Please check your API key in the .env file.",
            )
        elif "429" in str(e):
            raise HTTPException(
                status_code=500,
                detail="Rate limit exceeded. Please try again later.",
            )
        else:
            raise HTTPException(
                status_code=500,
                detail=f"Chat request failed: {str(e)}",
            )


@app.get("/")
def read_root():
    return {"message": "Welcome!"}


@app.get("/health")
def health_check():
    return {
        "status": "healthy",
        "client_ok": perplexity_client is not None
    }


class DownloadRequest(BaseModel):
    content: str





def create_docx_content(content: str) -> bytes:
    """Create a DOCX from the medical summary content"""
    doc = Document()
    
    # Add logo at the top right
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    try:
        # Add medical AI assistant branding
        logo_run = logo_paragraph.add_run("Medical AI Assistant")
        logo_run.bold = True
        logo_run.font.color.rgb = RGBColor(220, 53, 69)  # Red color
        logo_run.font.size = Pt(14)
    except Exception as e:
        # Fallback to text if formatting fails
        logo_run = logo_paragraph.add_run("Medical AI Assistant")
        logo_run.bold = True
        logo_run.font.color.rgb = RGBColor(220, 53, 69)  # Red color
        logo_run.font.size = Pt(14)
    
    # Add some space
    doc.add_paragraph()
    
    # Add title
    title = doc.add_heading('MEDICAL SUMMARY', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Clean content exactly like the frontend
    import re
    cleaned_content = content
    # Remove equal signs
    cleaned_content = re.sub(r'={3,}', '', cleaned_content)
    # Remove double asterisks
    cleaned_content = re.sub(r'\*\*', '', cleaned_content)
    # Remove single asterisks but keep content
    cleaned_content = re.sub(r'\*([^*]+)\*', r'\1', cleaned_content)
    # Remove standalone "MEDICAL SUMMARY" lines
    cleaned_content = re.sub(r'^MEDICAL SUMMARY\s*$', '', cleaned_content, flags=re.MULTILINE)
    # Remove any line containing "MEDICAL SUMMARY"
    cleaned_content = re.sub(r'^.*MEDICAL SUMMARY.*$', '', cleaned_content, flags=re.MULTILINE)
    # Remove all remaining instances of "MEDICAL SUMMARY"
    cleaned_content = re.sub(r'MEDICAL SUMMARY', '', cleaned_content)
    # Remove "HEADER SECTION" text
    cleaned_content = re.sub(r'HEADER SECTION', '', cleaned_content)
    
    # Process content sections
    sections = cleaned_content.split('###')
    
    for section in sections:
        if not section.strip():
            continue
            
        lines = section.strip().split('\n')
        if not lines:
            continue
            
        # First line is the section header
        header = lines[0].strip()
        if header:
            # Skip headers that are just formatting artifacts
            if header.lower() in ['header section', 'medical summary']:
                continue
            doc.add_heading(header, level=1)
        
        # Check if this section contains a table
        table_lines = [line for line in lines[1:] if '|' in line and line.strip()]
        
        if len(table_lines) >= 3:  # At least header, separator, and one data row
            # Process as table
            table_data = []
            for line in table_lines:
                if '---' not in line:  # Skip separator lines
                    cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                    if cells:
                        table_data.append(cells)
            
            if table_data:
                # Create table
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                
                # Fill table data
                for i, row_data in enumerate(table_data):
                    for j, cell_data in enumerate(row_data):
                        cell = table.cell(i, j)
                        cell.text = cell_data.replace('[', '').replace(']', '')
                        
                        # Style header row
                        if i == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
        else:
            # Process as regular text
            for line in lines[1:]:
                line = line.strip()
                if line and '|' not in line:  # Skip table lines
                    # Remove brackets from placeholders and clean up formatting
                    line = line.replace('[', '').replace(']', '').replace('**', '').replace('*', '')
                    if line:
                        doc.add_paragraph(line)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()





@app.post("/download-docx")
async def download_docx(request: DownloadRequest):
    try:
        docx_content = create_docx_content(request.content)
        return Response(
            content=docx_content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=medical_summary.docx"}
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")